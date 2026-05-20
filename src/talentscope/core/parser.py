"""
文档解析器
==========
支持 PDF / DOCX / TXT，统一返回纯文本。
"""
from __future__ import annotations

from pathlib import Path
from typing import IO, Union

FileLike = Union[str, Path, IO[bytes]]


def parse_document(file: FileLike, filename: str | None = None) -> str:
    """解析文档，返回纯文本。

    Args:
        file: 文件路径或 file-like 对象
        filename: 若 file 是 file-like，则需提供文件名以判断格式
    """
    if isinstance(file, (str, Path)):
        path = Path(file)
        ext = path.suffix.lower()
        with open(path, "rb") as f:
            return _dispatch(ext, f)
    else:
        if not filename:
            raise ValueError("file-like 输入必须同时提供 filename")
        ext = Path(filename).suffix.lower()
        return _dispatch(ext, file)


def _dispatch(ext: str, fh: IO[bytes]) -> str:
    if ext == ".pdf":
        return _parse_pdf(fh)
    if ext == ".docx":
        return _parse_docx(fh)
    if ext in (".txt", ".md"):
        return fh.read().decode("utf-8", errors="ignore")
    raise ValueError(f"不支持的文件格式: {ext}（支持 .pdf .docx .txt .md）")


def _parse_pdf(fh: IO[bytes]) -> str:
    from pypdf import PdfReader
    reader = PdfReader(fh)
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts).strip()


def _parse_docx(fh: IO[bytes]) -> str:
    from docx import Document
    # python-docx 接受 file-like
    doc = Document(fh)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 表格也提取
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts).strip()
